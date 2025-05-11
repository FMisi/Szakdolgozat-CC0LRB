from flask import Flask, jsonify, request, send_file
import json
import shutil
import requests
import subprocess
import logging
from urllib.parse import urlparse
from docx import Document
import os
from datetime import datetime
import zipfile
import io
from collections import Counter, defaultdict

app = Flask(__name__)

DEBUG = True

def initialize_log():
    with open("debug.log", "w") as log_file:
        pass

def log_debug_message(message):
    if DEBUG:
        with open("debug.log", "a") as log_file:
            log_file.write(message + "\n")
#region
@app.route('/')
def home():
    return open("templates/index.html").read()

@app.route('/scan', methods=['POST'])
def scan():
    data = request.get_json()
    url = data.get("url")
    mode = data.get("mode", "static")

    if not url:
        return jsonify({"hiba": "Nem adtál meg URL-t!"}), 400

    headers = {'User-Agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:136.0) Gecko/20100101 Firefox/136.0'}
    headers_str = f"{headers['User-Agent']}"
    
    # HTML letöltése egyszer, és használjuk az eszközök számára
    content = download_html(url, headers)
    if "error" in content:
        return jsonify(content), 500
#endregion
    log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] URL: {url} átadásra került a {mode.upper()} módban")

    # Eredmények az eszközök futtatása alapján
    scan_results = {}

    if mode == "static":
        scan_results["semgrep"] = run_semgrep(content)
    elif mode == "dynamic":
        scan_results["nikto"] = run_nikto(url)
        scan_results["arachni"] = run_arachni(url, headers_str)
        scan_results["wapiti"] = run_wapiti(url)
    elif mode == "llm":
        scan_results["llm"] = run_llm(content)
    elif mode == "all":
        scan_results["semgrep"] = run_semgrep(content)
        scan_results["nikto"] = run_nikto(url)
        scan_results["arachni"] = run_arachni(url, headers_str)
        scan_results["wapiti"] = run_wapiti(url)
        scan_results["llm"] = run_llm(content)
    else:
        return jsonify({"hiba": f"Ismeretlen mód: {mode}"}), 400

    combined_results = {
        "url": url,
        "scan_results": scan_results
    }

    log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] Teszteredmények sikeresen kombinálva.")
    return jsonify(combined_results)

def download_html(url, headers):
    try:
        print(f"Letöltjük a HTML-t az URL-ről: {url}...")

        # HTTP kérés a megadott URL-re, fejlécekkel
        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            print("HTML tartalom sikeresen letöltve.")
            return str(response.headers)+"\n\n"+str(response.text)  # Visszaadjuk a letöltött HTML tartalmat
        else:
            print(f"Nem sikerült letölteni a tartalmat. HTTP status: {response.status_code}")
            return {"error": f"Nem sikerült letölteni a tartalmat. HTTP status: {response.status_code}"}

    except requests.RequestException as e:
        print(f"Hiba történt a letöltés során: {e}")
        return {"error": f"Hiba történt a letöltés során: {str(e)}"}

#region run_semgrep
def run_semgrep(code):

    with open("response.html", "w") as f:
        f.write(code)

    result = subprocess.run([
        "semgrep", "--config=semgrep_rules/common_web_vulns.yml", "response.html", "--json", "-q"
    ], capture_output=True, text=True)

    print("Semgrep output:", result.stdout)
    
    try:
        output = json.loads(result.stdout)
        print("Parsed output:", output)
        
        return output
    
    except json.JSONDecodeError:
        print("Error occurred while running Semgrep.")
        return {"részletek": "Error decoding Semgrep output"}
#endregion

#region run_nikto
def run_nikto(url):
    try:
        print("Nikto fut...")

        # Nikto parancs beállítása
        nikto_command = [
            'nikto',
            '-h', url,
            '-maxtime', '90',
            '-Tuning', '1,2,3',
            '-nointeractive'
        ]

        # Nikto futtatása
        result = subprocess.run(
            nikto_command, text=True, capture_output=True, check=True
        )

        # Nikto kimenet feldolgozása
        if result.stdout:
            print("Nikto befejeződött...")
            return {"üzenet": "A scan sikeresen befejeződött!", "részletek": result.stdout}
        else:
            return {"hiba": "No output from Nikto"}

    except subprocess.CalledProcessError as e:
        logging.error(f"Hiba a Nikto futtatása közben: {str(e)}")
        return {"hiba": f"Hiba a Nikto futtatása közben: {str(e)}", "stderr": e.stderr}
    except Exception as e:
        logging.error(f"Váratlan hiba: {str(e)}")
        return {"hiba": f"Váratlan hiba: {str(e)}"}
#endregion

#region run_llm
def run_llm(response):
    try:
        print("Running LLM analysis with DeepSeek...")

        # Header -ök kivonata a válaszból
        headers = dict(response.headers)
        headers_str = json.dumps(headers, indent=2)

        # Prompt megadása a DeepSeek-nek az elemzéshez
        prompt = (
            "You are a security expert analyzing HTTP response headers. Below is a JSON representation of the headers from a web request:\n\n"
            f"{headers_str}\n\n"
            "Determine if the 'Strict-Transport-Security' header is present in the headers.\n"
            "If it is present, report its value and confirm it meets basic HSTS requirements (e.g., includes 'max-age').\n"
            "If it is not present, say 'Missing Strict-Transport-Security Header'. Explain the security implications of its absence if it's absent.\n"
            "Determine if the 'X-Cache' and 'Age' header is present in the headers.\n"
            "If both are present, say 'Possible Web Cache Poisoning Vector'. Explain the security implications of web cache poisoning if possible web cache poisoning vector is found.\n"
            "Provide your response in a clear, concise manner."
        )

        # A DeepSeek API függvény hívása a prompttal
        llm_response = get_deepseek_valasz(prompt)

        # A nyers LLM válasz naplózása debug -oláshoz
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] Raw LLM response: '{llm_response}'")

        if not llm_response or llm_response.strip() == "":
            return {
                "message": "LLM analysis failed",
                "error": "Empty response from DeepSeek API",
                "headers_analyzed": headers
            }
        elif llm_response.startswith("Error:") or llm_response.startswith("Exception:"):
            return {
                "message": "LLM analysis failed",
                "error": f"DeepSeek API error: {llm_response}",
                "headers_analyzed": headers
            }

        # Az LLM elemzésével visszatérés
        return {
            "message": "LLM analysis completed",
            "llm_response": llm_response,
            "headers_analyzed": headers
        }

    except Exception as e:
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] Exception in run_llm: {str(e)}")
        return {"error": f"Unexpected error in LLM analysis: {str(e)}", "headers_analyzed": headers}

def get_deepseek_valasz(prompt):
    API_KEY = "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"
    if not API_KEY:
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] API key not set")
        return "Error: API key not set"

    ENDPOINT = "https://openrouter.ai/api/v1/chat/completions"
    HEADERS = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    PAYLOAD = {
        "model": "deepseek/deepseek-chat:free",
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    try:
        # POST kérés küldése az API -nak
        response = requests.post(ENDPOINT, headers=HEADERS, data=json.dumps(PAYLOAD), timeout=10)
        response.raise_for_status()  # Kivétel dobása 4xx/5xx status code -ok esetében

        result = response.json()
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] API válasz: {json.dumps(result, indent=2)}")

        # Annak ellenőrzése, hogy az elvárt struktúrájú válasz létezik-e
        if "choices" in result and len(result["choices"]) > 0 and "message" in result["choices"][0]:
            return result["choices"][0]["message"]["content"]
        else:
            return "Hiba: Váratlan API válasz struktúra"

    except requests.exceptions.RequestException as e:
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] Kérés kivétel dobódott: {str(e)}")
        return f"Hiba: {response.status_code if response else 'Nincs válasz'}, {response.text if response else str(e)}"
    except Exception as e:
        log_debug_message(f"[{datetime.now().strftime('%Y-%m-%d %H:%M')}] A get_deepseek_valasz() -ban általános kivétel dobódott: {str(e)}")
        return f"Exception: {str(e)}"
#endregion

#region run_arachni
def run_arachni(url, headers):
    try:
        print("Arachni fut...")
        
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError("Érvénytelen URL: PKérlek adj meg egy teljes abszolút URL-t.")
        
        # Arachni riportok mentése egy külön mappába
        report_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'arachni_reports')
        if not os.path.exists(report_dir):
            os.makedirs(report_dir)
        
        report_file = os.path.join(report_dir, 'arachni_report.afr')
        
        arachni_command = [
            '/home/fmisi/Tools/arachni/bin/arachni', url,
            '--http-user-agent', str(headers),
            '--checks', '*',
            '--scope-page-limit', '2',
            '--scope-directory-depth-limit', '2',
            '--http-request-timeout', '1000',
            '--report-save-path', report_file
        ]
        
        result = subprocess.run(
            arachni_command, text=True, capture_output=True, check=True
        )
        
        if result.returncode != 0:
            return {"hiba": "Arachni hiba", "részletek": result.stderr}
        
        output_lines = result.stdout.splitlines()
        start_index = next((i for i, line in enumerate(output_lines) if "[~] ===========================" in line), None)
        filtered_output = "\n".join(output_lines[start_index:]) if start_index is not None else "Nincs releváns adat az Arachni kimenetében."
        
        # A jelentés generálása után töröljük az .afr fájlokat
        shutil.rmtree(report_dir)
        
        return {"üzenet": "Arachni scan befejeződött", "részletek": filtered_output}
    
    except subprocess.CalledProcessError as e:
        return {"hiba": f"Arachni hiba: {str(e)}", "stderr": e.stderr}
    except Exception as e:
        return {"hiba": f"Váratlan hiba: {str(e)}"}
#endregion

#region run_wapiti
def run_wapiti(url):
    try:
        print("Wapiti futtatása...")

        # Ellenőrizzük, hogy érvényes URL-t kaptunk
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError("Érvénytelen URL: Kérjük, adjon meg egy teljes, abszolút URL-t.")

        # A program szkript melletti /temp mappa elérési útja
        temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'temp')

        # Ha még nem létezik a temp mappa, létrehozzuk
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        # Kimeneti fájl
        result_file = os.path.join(temp_dir, "wapiti_results.json")

        # Wapiti parancs összeállítása az adott URL-lel
        wapiti_command = [
            'wapiti', '-u', url,
            '--header', 'User-Agent: Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:136.0) Gecko/20100101 Firefox/136.0',
            '--timeout', '1',
            '--max-links-per-page', '1',
            '--max-files-per-dir', '1',
            '--max-scan-time', '60',
            '--scan-force', 'paranoid',  # paranoid, sneaky, polite, normal, aggressive, insane
            '--format', 'json',
            '--output', result_file
        ]

        # Wapiti futtatása
        result = subprocess.run(
            wapiti_command, text=True, capture_output=True, check=True
        )

        # Wapiti kimenet ellenőrzése
        if result.returncode == 0:
            print("Wapiti befejeződött.")
            
            # A JSON fájl olvasása
            with open(result_file, 'r') as report_file:
                report_content = json.load(report_file)
            
            # Kivágjuk a "vulnerabilities" részt
            if 'vulnerabilities' in report_content:
                vulnerabilities = report_content['vulnerabilities']

                # A jelentés generálása után töröljük az ideiglenes .json fájlokat
                shutil.rmtree(temp_dir)

                return {"üzenet": "A vizsgálat sikeresen befejeződött", "vulnerabilities": vulnerabilities}
            else:
                return {"hiba": "Nincs sebezhetőség a vizsgálatban"}
        else:
            return {"hiba": "A Wapiti nem adott vissza kimenetet"}

    except subprocess.CalledProcessError as e:
        logging.error(f"Hiba a Wapiti futtatása közben: {str(e)}")
        return {"hiba": f"Hiba a Wapiti futtatása közben: {str(e)}", "stderr": e.stderr}
    except ValueError as e:
        logging.error(f"Érvénytelen URL: {str(e)}")
        return {"hiba": f"Érvénytelen URL: {str(e)}"}
    except Exception as e:
        logging.error(f"Váratlan hiba: {str(e)}")
        return {"hiba": f"Váratlan hiba: {str(e)}"}
#endregion

#region generate_report
@app.route('/generate_report', methods=['POST'])
def generate_report():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"hiba": "Nincs adat megadva a jelentésgeneráláshoz"}), 400
        
        url = data.get("url", "N/A")
        scan_results = data.get("scan_results", {})
        
        # Sérülékenységek számlálása és eszközök követése
        vuln_counter = Counter()  # Előfordulások száma
        vuln_tools = defaultdict(set)  # Eszközök halmaza sérülékenységenként
        
        for tool, results in scan_results.items():
            #region semgrep
            if tool == "semgrep" and "results" in results:
                details = results["results"]

                csrf_occurences = sum(1 for result in details if result.get("check_id") == "semgrep_rules.missing-csrf-token")

                if csrf_occurences > 0:
                    vuln_counter["Cross-site Request Forgery (CSRF)"] += csrf_occurences
                    vuln_tools["Cross-site Request Forgery (CSRF)"].add(tool)
            #endregion
            #region nikto
            elif tool == "nikto" and "részletek" in results:
                details = results["részletek"]
                #region nikto_occurences
                hsts_occurrences = details.count("The site uses TLS and the Strict-Transport-Security HTTP header is not defined.")
                mimetype_occurences = details.count("The X-Content-Type-Options header is not set")
                csp_occurences = details.count("X-Frame-Options header is deprecated")
                no_secure_flag_occurences = details.count("without the secure flag")
                no_httponly_flag_occurences = details.count("without the httponly flag")
                #endregion
                if hsts_occurrences > 0:
                    vuln_counter["HTTP Strict Transport Security (HSTS)"] += hsts_occurrences
                    vuln_tools["HTTP Strict Transport Security (HSTS)"].add(tool)
                if mimetype_occurences > 0:
                    vuln_counter["MIME Type Confusion"] += mimetype_occurences
                    vuln_tools["MIME Type Confusion"].add(tool)
                if csp_occurences > 0:
                    vuln_counter["Content Security Policy Misconfiguration"] += csp_occurences
                    vuln_tools["Content Security Policy Misconfiguration"].add(tool)
                if "Server: No banner retrieved" not in details:
                    vuln_counter["Information Disclosure (Server Banner)"] += 1
                    vuln_tools["Information Disclosure (Server Banner)"].add(tool)
                if no_secure_flag_occurences > 0:
                    vuln_counter["Session Cookie Without the secure Flag Set"] += no_secure_flag_occurences
                    vuln_tools["Session Cookie Without the secure Flag Set"].add(tool)
                if no_httponly_flag_occurences > 0:
                    vuln_counter["Session Cookie Without the httponly Flag Set"] += no_httponly_flag_occurences
                    vuln_tools["Session Cookie Without the httponly Flag Set"].add(tool)
            #endregion
            #region arachni
            elif tool == "arachni" and "részletek" in results:
                details = results["részletek"]
                
                # Biztosítjuk, hogy a kulcsok léteznek
                for vuln in ["HTTP Strict Transport Security (HSTS)", "Common Directory Found", "Interesting Response", "Blind SQL Injection (differential analysis)", "Blind SQL Injection (timing attack)", "Cross-Site Scripting (XSS)", "Cross-Site Request Forgery"]:
                    if vuln not in vuln_counter:
                        vuln_counter[vuln] = 0
                    if vuln not in vuln_tools:
                        vuln_tools[vuln] = set()

                if isinstance(details, str):
                    hsts_count = details.count("Missing 'Strict-Transport-Security' header")
                    if hsts_count > 0:
                        vuln_counter["HTTP Strict Transport Security (HSTS)"] += hsts_count
                        vuln_tools["HTTP Strict Transport Security (HSTS)"].add(tool)

                    common_dir_count = details.count("Common directory")
                    if common_dir_count > 0:
                        vuln_counter["Common Directory Found"] += common_dir_count
                        vuln_tools["Common Directory Found"].add(tool)

                    interesting_count = details.count("Interesting response")
                    if interesting_count > 0:
                        vuln_counter["Interesting Response"] += interesting_count
                        vuln_tools["Interesting Response"].add(tool)

                    blind_sqli_diff_analysis_count = details.count("Blind SQL Injection (differential analysis)")
                    if blind_sqli_diff_analysis_count > 0:
                        vuln_counter["Blind SQL Injection (differential analysis)"] += blind_sqli_diff_analysis_count
                        vuln_tools["Blind SQL Injection (differential analysis)"].add(tool)

                    blind_sqli_timing_count = details.count("Blind SQL Injection (timing attack)")
                    if blind_sqli_timing_count > 0:
                        vuln_counter["Blind SQL Injection (timing attack)"] += blind_sqli_timing_count
                        vuln_tools["Blind SQL Injection (timing attack)"].add(tool)

                    xss_count = details.count("Cross-Site Scripting (XSS)")
                    if xss_count > 0:
                        vuln_counter["Cross-Site Scripting (XSS)"] += xss_count
                        vuln_tools["Cross-Site Scripting (XSS)"].add(tool)

                    csrf_occurences = details.count("Cross-Site Request Forgery")
                    if csrf_occurences > 0:
                        vuln_counter["Cross-Site Request Forgery"] += csrf_occurences
                        vuln_tools["Cross-Site Request Forgery"].add(tool)

                elif isinstance(details, list):
                    for item in details:
                        if "Missing 'Strict-Transport-Security' header" in item:
                            vuln_counter["Missing HTTP Strict Transport Security (HSTS)"] += 1
                            vuln_tools["Missing HTTP Strict Transport Security (HSTS)"].add(tool)
                        if "Common directory" in item:
                            vuln_counter["Common Directory Found"] += 1
                            vuln_tools["Common Directory Found"].add(tool)
                        if "Interesting response" in item:
                            vuln_counter["Interesting Response"] += 1
                            vuln_tools["Interesting Response"].add(tool)
                        if "Blind SQL Injection (differential analysis)" in item:
                            vuln_counter["Blind SQL Injection (differential analysis)"] += 1
                            vuln_tools["Blind SQL Injection (differential analysis)"].add(tool)
                        if "Blind SQL Injection (timing attack)" in item:
                            vuln_counter["Blind SQL Injection (timing attack)"] += 1
                            vuln_tools["Blind SQL Injection (timing attack)"].add(tool)
                        if "Cross-Site Scripting (XSS)" in item:
                            vuln_counter["Cross-Site Scripting (XSS)"] += 1
                            vuln_tools["Cross-Site Scripting (XSS)"].add(tool)
            #endregion
            #region wapiti
            elif tool == "wapiti" and "vulnerabilities" in results:
                for vuln_type, vuln_list in results["vulnerabilities"].items():
                    if vuln_list:
                        vuln_counter[vuln_type] += len(vuln_list)
                        vuln_tools[vuln_type].add(tool)
            #endregion
            #region llm
            elif tool == "llm" and "llm_response" in results:
                llm_response = results["llm_response"]
                hsts_occurrences = llm_response.count("Missing Strict-Transport-Security Header")
                web_cache_poisoning_occurences = llm_response.count("Possible Web Cache Poisoning Vector")
                if hsts_occurrences > 0:
                    vuln_counter["HTTP Strict Transport Security (HSTS)"] += hsts_occurrences
                    vuln_tools["HTTP Strict Transport Security (HSTS)"].add(tool)
                if web_cache_poisoning_occurences > 0:
                    vuln_counter["Possible Web Cache Poisoning Vector"] += web_cache_poisoning_occurences
                    vuln_tools["Possible Web Cache Poisoning Vector"].add(tool)
            #endregion
        # Sérülékenységek csökkenő sorrendbe rendezése az előfordulások száma alapján
        sorted_vulns = sorted(vuln_counter.items(), key=lambda x: x[1], reverse=True)

        # Riport generálása
        doc = Document()
        doc.add_heading('Biztonsági Vizsgálati Jelentés', level=1)
        doc.add_paragraph(f"Vizsgált URL: {url}")
        
        for vuln, count in sorted_vulns:
            if count > 0:
                tool_count = len(vuln_tools[vuln])
                doc.add_heading(f"{vuln} ({count})", level=2)
                doc.add_paragraph(f"Detektálva {tool_count} eszköz által.")
        
        # Jelentés mentése
        report_file = "vizsgalat_report.docx"
        doc.save(report_file)
        
        return send_file(report_file, as_attachment=True)
    
    except Exception as e:
        return jsonify({"hiba": f"Hiba a jelentés generálása során: {str(e)}"}), 500
#endregion

#region download_raw_output
@app.route('/download_raw_zip', methods=['POST'])
def download_raw_zip():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"hiba": "Nincs beérkező adat a ZIP generálásához"}), 400
        
        scan_results = data.get("scan_results", {})
        if not scan_results:
            return jsonify({"hiba": "Nem találhatóak vizsgálati eredmények"}), 400
        
        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            for tool, result in scan_results.items():
                if not result:
                    continue  # Üres eredményt nem mentünk el
                
                filename = f"{tool}.json"
                
                # JSON formázás grep-barát módon
                if isinstance(result, (dict, list)):
                    content = json.dumps(result, indent=4, ensure_ascii=False).replace("\\n", "\n")
                else:
                    content = str(result).replace("\\n", "\n")  # Ha a kimenetben \n van, akkor az valódi új sor legyen
                
                zf.writestr(filename, content)
        
        if not zf.filelist:  # Ha nincs mit menteni, hibaüzenetet adunk
            return jsonify({"hiba": "A nyers output fájl üres, nincs mit letölteni!"}), 400

        memory_file.seek(0)
        return send_file(memory_file, download_name="raw_outputs.zip", as_attachment=True)
    
    except Exception as e:
        return jsonify({"hiba": f"Hiba a ZIP generálása közben: {str(e)}"}), 500
#endregion



#region download_log
@app.route('/download_log')
def download_log():
    log_file_path = "debug.log"
    if os.path.exists(log_file_path):
        return send_file(log_file_path, as_attachment=True)
    else:
        return jsonify({"hiba": "Nincs elérhető log fájl!"}), 404
#endregion

if __name__ == '__main__':
    initialize_log()
    app.run(debug=True)
